"""Generate MatchXP Session Report as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    # Light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_note(doc, label, text):
    p = doc.add_paragraph()
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    p.add_run(text)
    return p


def add_screenshot_note(doc, where):
    p = doc.add_paragraph()
    run = p.add_run(f"📸 Screenshot: {where}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x44, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("MatchXP — Bug Fix Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Sessions: 2026-05-18 · 2026-05-19 · 2026-05-20").alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)
doc.add_paragraph(
    "Project: Flutter Dating App with Game-Gated Chat | Backend: Supabase"
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Overview ───────────────────────────────────────────────────────────────────
add_heading(doc, "1. Overview", level=1)
doc.add_paragraph(
    "This report documents 26 fixes applied across three development sessions "
    "(May 18–20 2026) to the MatchXP Flutter dating app. The fixes span database "
    "migrations, Row-Level Security (RLS) policy corrections, Flutter game-screen "
    "logic, chat screen redesign, and UX improvements."
)
doc.add_paragraph()

# ── Session 2026-05-18 ─────────────────────────────────────────────────────────
add_heading(doc, "2. Session 2026-05-18", level=1)

# FIX 1
add_heading(doc, "FIX 1 — Emoji Charades: session never dying (re-entry loop)", level=2)
add_note(doc, "File", "lib/games/emoji_charades/emoji_charades_game_screen.dart")
doc.add_paragraph(
    "After finishing Emoji Charades, opening the same chat auto-navigated back into "
    "the game every time. Word Search and RPS were unaffected."
)
doc.add_paragraph("Root cause (multi-layer):")
doc.add_paragraph(
    "_onPartnerUpdate() only acted when _phase == waitSolve. If the partner's "
    "realtime event arrived during solve phase, it was silently dropped.",
    style="List Bullet",
)
doc.add_paragraph(
    "The amDesignated system (only the winner calls the RPC) was computed from "
    "_outcome, which depends on stale _partnerRow. Wrong designation → RPC never called.",
    style="List Bullet",
)
doc.add_paragraph(
    "_resumePendingChallengePoller found the active session and re-navigated on every "
    "chat open.",
    style="List Bullet",
)
doc.add_paragraph("Fixes applied:")
doc.add_paragraph(
    "_checkCompletionOnLoad() — called after every initial record fetch; if both "
    "players already finished, jumps to _Phase.done and calls _sendCompletion().",
    style="List Bullet",
)
doc.add_paragraph(
    "_startCompletionPoller() — Timer that polls emoji_charades_games every 3 s in "
    "waitSolve phase; calls _onPartnerUpdate() when partner finishes.",
    style="List Bullet",
)
doc.add_paragraph(
    "_sendCompletion() simplified — removed amDesignated logic; both players call "
    "complete_game_session unconditionally. DB-level idempotency handles duplicates.",
    style="List Bullet",
)
add_screenshot_note(doc, "Open Emoji Charades game screen → show the waitSolve phase UI")

# FIX 2
add_heading(doc, "FIX 2 — Session expiry (5 minutes)", level=2)
add_note(doc, "File", "lib/screens/chat_conversation_screen.dart")
doc.add_paragraph(
    "Expired/stuck game invitations kept showing 'Game is in progress!' and the resume "
    "poller would auto-navigate users to old sessions."
)
doc.add_paragraph("Changes:")
doc.add_paragraph("_kInviteExpiry reduced to Duration(minutes: 5).", style="List Bullet")
doc.add_paragraph(
    "Expiry now covers both pending AND active sessions.", style="List Bullet"
)
doc.add_paragraph(
    "Expired active sessions show 'Game session expired'.", style="List Bullet"
)
doc.add_paragraph(
    "_resumePendingChallengePoller skips sessions older than 5 minutes.",
    style="List Bullet",
)
add_screenshot_note(
    doc,
    "Chat screen — game invitation card showing 'Game session expired' state (dark-grey card)",
)

# FIX 3
add_heading(
    doc,
    "FIX 3 — complete_game_session made idempotent (DB migration: idempotent_complete_game_session)",
    level=2,
)
doc.add_paragraph(
    "With both players calling complete_game_session, duplicate game_result messages "
    "were being inserted in the chat."
)
doc.add_paragraph("Migration SQL:")
add_code_block(
    doc,
    """UPDATE game_sessions
SET status = 'completed', completed_at = NOW()
WHERE id = p_session_id AND status != 'completed';

GET DIAGNOSTICS v_rows_updated = ROW_COUNT;
IF v_rows_updated = 0 THEN RETURN; END IF;  -- already completed, skip INSERT""",
)
doc.add_paragraph(
    "Only the first concurrent call proceeds to insert the game_result bubble. "
    "Any subsequent call finds status = 'completed' and returns early."
)

# FIX 4
add_heading(
    doc,
    "FIX 4 — session_id never written to emoji_charades_games",
    level=2,
)
add_note(doc, "File", "lib/games/emoji_charades/emoji_charades_game_screen.dart")
doc.add_paragraph(
    "_upsertMyRecord() never included session_id. Every row had session_id = null, "
    "so queries on the same match returned rows from previous games."
)
doc.add_paragraph("Fixes:")
doc.add_paragraph(
    "_upsertMyRecord() now includes session_id on every write.", style="List Bullet"
)
doc.add_paragraph(
    "_loadMyRecord(), initial partner fetch in _subscribeRealtime(), and "
    "_startCompletionPoller() all add .eq('session_id', widget.sessionId!) when "
    "sessionId is non-null.",
    style="List Bullet",
)
doc.add_paragraph(
    "Realtime callback ignores events where row['session_id'] != widget.sessionId.",
    style="List Bullet",
)

# FIX 5
add_heading(
    doc,
    "FIX 5 — RPS: first player stuck on waiting screen forever (DB migration: fix_rps_moves_rls_circular_dependency)",
    level=2,
)
doc.add_paragraph(
    "Player A (who submitted first) was permanently stuck on the waiting screen. "
    "All detection paths (realtime, subscribeToSession fallback, 3-second poll) failed."
)
doc.add_paragraph("Root cause — circular RLS policy:")
add_code_block(
    doc,
    """-- OLD (broken)
(auth.uid() = player_id) OR
(SELECT count(*) FROM rps_moves WHERE session_id = ...) >= 2
-- PostgreSQL evaluates the count under the same RLS context → count never reaches 2""",
)
doc.add_paragraph("DB fix:")
add_code_block(
    doc,
    """DROP POLICY rps_select ON rps_moves;
CREATE POLICY rps_select ON rps_moves FOR SELECT USING (
  EXISTS (
    SELECT 1 FROM game_sessions
    WHERE  id             = rps_moves.session_id
      AND  (challenger_id = auth.uid() OR challenged_id = auth.uid())
  )
);""",
)
add_screenshot_note(
    doc, "RPS waiting screen — show both players having submitted and seeing the reveal"
)

# FIX 6
add_heading(
    doc,
    "FIX 6 — RPS: chat card stays 'Game is in progress!' after finishing",
    level=2,
)
add_note(doc, "File", "lib/screens/chat_conversation_screen.dart")
doc.add_paragraph(
    "After finishing RPS and returning to chat, the invitation card still showed "
    "'Game is in progress!'. RPS uses pushReplacement at each step, so the .then(refresh) "
    "callback fired immediately after RPSPickScreen was replaced — not after the game ended."
)
doc.add_paragraph("Fix — added refresh calls inside the RPS onChatUnlocked callback:")
add_code_block(
    doc,
    """// lib/screens/chat_conversation_screen.dart — _navigateToGame, rps case
onChatUnlocked: () {
  if (!mounted) return;
  setState(() => _chatUnlocked = true);
  _restartMessageStream();   // added
  _refreshSessionStream();   // added
},""",
)
add_screenshot_note(
    doc, "Chat screen — RPS game card showing 'Game completed ✓' after finishing"
)

# FIX 7
add_heading(
    doc,
    "FIX 7 — Black screen after Google Sign-In back navigation (Impeller disabled)",
    level=2,
)
add_note(doc, "File", "android/app/src/main/AndroidManifest.xml")
doc.add_paragraph(
    "Navigating back to the login screen after a Google Sign-In attempt caused a "
    "black screen on Android emulators. Flutter's Impeller rendering engine conflicts "
    "with platform views used by the Google Sign-In overlay."
)
doc.add_paragraph("Fix:")
add_code_block(
    doc,
    """<!-- Disable Impeller — fixes black screen on emulators -->
<meta-data
    android:name="io.flutter.embedding.android.EnableImpeller"
    android:value="false" />""",
)
add_screenshot_note(
    doc, "Login screen — shown after returning from Google Sign-In (no black screen)"
)

doc.add_paragraph()

# ── Session 2026-05-19 ─────────────────────────────────────────────────────────
add_heading(doc, "3. Session 2026-05-19", level=1)

# FIX 8
add_heading(doc, "FIX 8 — Word Search: lint cleanup in wrapper", level=2)
add_note(
    doc, "File", "lib/games/word_search/word_search_supabase_wrapper.dart"
)
doc.add_paragraph(
    "Removed unused import package:google_fonts/google_fonts.dart and unused field "
    "bool _loading after an abandoned spinner approach."
)

# FIX 9
add_heading(
    doc,
    "FIX 9 — RPS & Word Search: intro/welcome screens were being skipped",
    level=2,
)
add_note(
    doc,
    "Files",
    "lib/games/rock_paper_scissors/screens/rps_intro_screen.dart, "
    "lib/screens/chat_conversation_screen.dart, lib/games/game_hub_screen.dart, "
    "lib/games/word_search/word_search_supabase_wrapper.dart",
)
doc.add_paragraph(
    "Both games have welcome/rules screens but players were dropped straight into "
    "gameplay. RPSIntroScreen existed but was never pushed. Word Search passed "
    "skipWelcome: true."
)
doc.add_paragraph("Fixes:")
doc.add_paragraph(
    "Added popCount param to RPSIntroScreen so it threads through to RPSResultScreen._startChat().",
    style="List Bullet",
)
doc.add_paragraph(
    "Both chat_conversation_screen.dart and game_hub_screen.dart now push RPSIntroScreen.",
    style="List Bullet",
)
doc.add_paragraph(
    "Changed skipWelcome: true → skipWelcome: false in WordSearchSupabaseWrapper.build().",
    style="List Bullet",
)
add_screenshot_note(
    doc,
    "RPS intro/rules screen and Word Search welcome screen — first screen users see before gameplay",
)

# FIX 10
add_heading(doc, "FIX 10 — Word Search: win/loss/draw result logic rewritten", level=2)
add_note(
    doc,
    "Files",
    "lib/games/word_search/word_search_models.dart, word_search_service.dart, word_search_supabase_wrapper.dart",
)
doc.add_paragraph("New result rules:")
doc.add_paragraph(
    "You solve, partner skips → you win", style="List Bullet"
)
doc.add_paragraph(
    "You skip, partner solves → you lose", style="List Bullet"
)
doc.add_paragraph(
    "Both solve → draw", style="List Bullet"
)
doc.add_paragraph(
    "Both skip → draw", style="List Bullet"
)
doc.add_paragraph(
    "Added skipped to GameStatus enum; added markSkipped() to word_search_service.dart; "
    "win/draw logic now based on isSkipped flags, not timestamps."
)
add_screenshot_note(
    doc, "Word Search done/result screen — showing win, lose, or draw label"
)

# FIX 11
add_heading(
    doc,
    "FIX 11 — Word Search: submit and solve/skip always show intermediate waiting screens",
    level=2,
)
add_note(doc, "File", "lib/games/word_search/word_search_game_screen.dart")
doc.add_paragraph(
    "If the partner had already submitted or solved, the current player's submit/solve "
    "jumped straight past the waiting screen without any confirmation."
)
doc.add_paragraph("Fixes:")
doc.add_paragraph(
    "_submitPuzzle() always sets step = WSStep.wait. An 800 ms timer auto-advances if "
    "partner already submitted.",
    style="List Bullet",
)
doc.add_paragraph(
    "Added _goToWaitPartner(): always sets step = WSStep.waitPartner. If partner is "
    "already done, 800 ms timer auto-advances to done.",
    style="List Bullet",
)
add_screenshot_note(
    doc, "Word Search wait screen ('Waiting for partner...') and waitPartner screen"
)

# FIX 12
add_heading(
    doc,
    "FIX 12 — Word Search: session isolation (DB migration: add_session_id_to_word_search_games)",
    level=2,
)
doc.add_paragraph(
    "Playing Word Search a second time on the same match showed session 1's solved "
    "puzzle. getSnapshot queried by match_id only with no ORDER BY."
)
doc.add_paragraph("DB migration SQL:")
add_code_block(
    doc,
    """ALTER TABLE word_search_games
  ADD COLUMN IF NOT EXISTS session_id UUID REFERENCES game_sessions(id);
CREATE INDEX IF NOT EXISTS idx_wsg_session_id ON word_search_games(session_id);""",
)
doc.add_paragraph("Flutter fixes: session_id added to model, service, and wrapper INSERT.")

# FIX 13
add_heading(
    doc,
    "FIX 13 — Emoji Charades: stale solved/skipped flags from previous session",
    level=2,
)
add_note(doc, "File", "lib/games/emoji_charades/emoji_charades_game_screen.dart")
doc.add_paragraph(
    "UNIQUE (match_id, user_id) constraint means one row per user per match. "
    "_sendPuzzle() upsert did not reset solved/skipped, so session 1 flags persisted "
    "into session 2, causing premature game-over."
)
doc.add_paragraph("Fix — _sendPuzzle() upsert now includes:")
add_code_block(
    doc,
    """'solved': false,   // reset stale flags from any previous session on this match
'skipped': false,""",
)

# FIX 14
add_heading(
    doc,
    "FIX 14 — Emoji Charades: skips waiting screen when partner already submitted",
    level=2,
)
add_note(doc, "File", "lib/games/emoji_charades/emoji_charades_game_screen.dart")
doc.add_paragraph(
    "If User B submitted before User A, User A would skip the waiting screen and land "
    "directly on the solve screen."
)
doc.add_paragraph(
    "Fix: _sendPuzzle() now always calls _goTo(_Phase.waiting). "
    "The waiting screen renders 'Solve Partner's Puzzle' button when _partnerSubmitted = true."
)
add_screenshot_note(
    doc, "Emoji Charades waiting screen — 'Solve [Partner]'s Puzzle' button visible"
)

# FIX 15
add_heading(
    doc,
    "FIX 15 — Word Search: auto-advances to solve screen without user input",
    level=2,
)
add_note(doc, "File", "lib/games/word_search/word_search_game_screen.dart")
doc.add_paragraph(
    "After submitting, users were automatically redirected to the solve screen in two "
    "cases without any tap (800 ms timer and setPartnerData state update)."
)
doc.add_paragraph("Fixes:")
doc.add_paragraph(
    "Removed 800 ms auto-advance block from _submitPuzzle().", style="List Bullet"
)
doc.add_paragraph(
    "Removed auto-advance from setPartnerData().", style="List Bullet"
)
doc.add_paragraph(
    "_buildWait() now shows 'Solve [Partner]'s Puzzle' button when partner.submitted = true.",
    style="List Bullet",
)
add_screenshot_note(
    doc,
    "Word Search wait screen showing '[Partner]'s puzzle is ready!' with action button",
)

doc.add_paragraph()

# ── Session 2026-05-20 ─────────────────────────────────────────────────────────
add_heading(doc, "4. Session 2026-05-20", level=1)

# FIX 19
add_heading(
    doc,
    "FIX 19 — Removed game/controller option from '+' media picker in chat",
    level=2,
)
add_note(doc, "File", "lib/screens/chat_conversation_screen.dart")
doc.add_paragraph(
    "Deleted the controller tile from _showMediaPicker() bottom sheet. "
    "Game challenges are now sent exclusively via the controller button in the app bar."
)
add_screenshot_note(
    doc,
    "Chat screen '+' bottom sheet — controller tile is no longer present",
)

# FIX 20
add_heading(
    doc,
    "FIX 20 — Controller button in chat app bar now sends game challenge via GameHub",
    level=2,
)
add_note(
    doc, "Files", "lib/games/game_hub_screen.dart, lib/screens/chat_conversation_screen.dart"
)
doc.add_paragraph(
    "Previously, tapping the controller icon opened GameHub but PLAY launched the game "
    "locally — the partner received no invitation."
)
doc.add_paragraph("Fix — two parts:")
doc.add_paragraph(
    "GameHubScreen: added optional onGameSelected callback. Each PLAY button checks "
    "onGameSelected != null → pops screen and calls callback instead of launching locally.",
    style="List Bullet",
)
doc.add_paragraph(
    "Controller button in chat_conversation_screen.dart passes "
    "onGameSelected: _sendGameChallenge to GameHub, which creates the challenge via "
    "create_game_challenge RPC and broadcasts to the partner.",
    style="List Bullet",
)
add_screenshot_note(
    doc,
    "Chat screen → controller button → GameHub → PLAY → challenge card appears in chat for both users",
)

# FIX 21
add_heading(
    doc, "FIX 21 — GameHub title and subtitle gradient matches PLAY button", level=2
)
add_note(doc, "File", "lib/games/game_hub_screen.dart")
doc.add_paragraph(
    "Both 'Pick a game to play' and the subtitle now use ShaderMask with "
    "LinearGradient(colors: [Color(0xFF7C3AED), Color(0xFFA78BFA)]) — matching the PLAY button."
)
add_screenshot_note(doc, "GameHub screen — title and subtitle gradient style")

# FIX 22
add_heading(
    doc,
    "FIX 22 — 'It's a Match!' popup not shown to User A (first to like)",
    level=2,
)
add_note(doc, "File", "lib/screens/main_navigation.dart")
doc.add_paragraph(
    "Only User B (second to like) ever saw the match popup. User A's likeProfile() "
    "returned null (User B hadn't liked back yet) so no popup was triggered."
)
doc.add_paragraph(
    "Fix — 15-second polling in MainNavigation (always mounted):"
)
doc.add_paragraph(
    "initState seeds _knownMatchIds from the initial matches load.",
    style="List Bullet",
)
doc.add_paragraph(
    "Timer.periodic every 15 s calls _pollForNewMatches(). New matches → banner: "
    "'It's a Match! 🎉 You and [name] liked each other! Tap to play.'",
    style="List Bullet",
)
doc.add_paragraph(
    "Tapping the banner opens GameHub for that match.", style="List Bullet"
)
add_screenshot_note(
    doc,
    "MainNavigation notification banner — 'It's a Match!' shown to User A within 15 seconds",
)

# FIX 23
add_heading(
    doc,
    "FIX 23 — Match popup button: 'Play Game' → 'PLAY GAME TO UNLOCK CHAT'",
    level=2,
)
add_note(doc, "File", "lib/screens/home_screen.dart")
doc.add_paragraph(
    "Updated the primary button label in the HomeScreen 'It's a Match!' popup to match "
    "the same label used in ChatsScreen (FIX 17)."
)
add_screenshot_note(
    doc, "HomeScreen — 'It's a Match!' popup with updated button label"
)

# FIX 24
add_heading(
    doc,
    "FIX 24 — Chats screen redesigned: 'Pending Games' circles + split Messages section",
    level=2,
)
add_note(doc, "File", "lib/screens/chats_screen.dart")
doc.add_paragraph(
    "Top circle row now shows only locked matches with color-coded borders:"
)
doc.add_paragraph("Grey — no challenge sent yet", style="List Bullet")
doc.add_paragraph("Green — I sent the challenge", style="List Bullet")
doc.add_paragraph("Red — they sent me the challenge", style="List Bullet")
doc.add_paragraph(
    "Purple — both submitted puzzles (solving phase)", style="List Bullet"
)
doc.add_paragraph(
    "Messages section now shows only unlocked matches. "
    "Empty state: 'Play a game to start chatting!' when no unlocked matches exist."
)
doc.add_paragraph(
    "Section hidden entirely when no locked matches exist — Messages fills full height."
)
add_screenshot_note(
    doc,
    "Chats screen — Pending Games circle row (grey/green/red/purple borders) + Messages section below",
)

# FIX 25
add_heading(
    doc,
    "FIX 25 — Removed game controller button from chats screen header",
    level=2,
)
add_note(doc, "File", "lib/screens/chats_screen.dart")
doc.add_paragraph(
    "Removed the sports_esports_rounded button from the chats header. "
    "Also removed _showGameStatusPage(), game_status_screen.dart import, "
    "and match popup widget code."
)
add_screenshot_note(doc, "Chats screen header — controller icon is no longer present")

# FIX 26
add_heading(
    doc,
    "FIX 26 — 'Rejoin Game' button on active game challenge cards + disabled auto-navigation",
    level=2,
)
add_note(doc, "File", "lib/screens/chat_conversation_screen.dart")
doc.add_paragraph(
    "After navigating away from an active session there was no way to re-enter. "
    "_resumePendingChallengePoller auto-navigated but felt aggressive."
)
doc.add_paragraph("Fix — two parts:")
doc.add_paragraph(
    "Added canRejoin flag (liveStatus == 'active' && sessionId != null && !isExpired) "
    "to _buildGameChallengeBubble. When true, a purple 'Rejoin Game' button is shown.",
    style="List Bullet",
)
doc.add_paragraph(
    "Commented out _resumePendingChallengePoller() call in initState — auto-navigation "
    "is disabled but the method is kept for future use.",
    style="List Bullet",
)
add_code_block(
    doc,
    "// _resumePendingChallengePoller(); // disabled — use 'Rejoin Game' button instead",
)
add_screenshot_note(
    doc,
    "Chat screen — active game challenge card with purple 'Rejoin Game' button",
)

doc.add_paragraph()

# ── Summary Table ──────────────────────────────────────────────────────────────
add_heading(doc, "5. Summary of DB Migrations", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Migration name"
hdr[1].text = "Purpose"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ("idempotent_complete_game_session", "Atomic guard on complete_game_session RPC — first caller wins, subsequent calls are no-ops"),
    ("fix_rps_moves_rls_circular_dependency", "Replace circular count-based rps_select RLS policy with a clean game_sessions participant check"),
    ("add_session_id_to_word_search_games", "Add session_id column + index to word_search_games for per-session isolation"),
]
for name, purpose in rows:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = purpose

doc.add_paragraph()

add_heading(doc, "6. Summary of Files Modified", level=1)
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "File"
hdr2[1].text = "Session"
hdr2[2].text = "Fixes"
for cell in hdr2:
    for run in cell.paragraphs[0].runs:
        run.bold = True

file_rows = [
    ("lib/games/emoji_charades/emoji_charades_game_screen.dart", "05-18, 05-19", "1, 4, 13, 14"),
    ("lib/screens/chat_conversation_screen.dart", "05-18, 05-19, 05-20", "2, 6, 9, 18, 19, 20, 26"),
    ("android/app/src/main/AndroidManifest.xml", "05-18", "7"),
    ("lib/games/word_search/word_search_supabase_wrapper.dart", "05-19", "8, 10, 12"),
    ("lib/games/word_search/word_search_game_screen.dart", "05-19", "9, 11, 15"),
    ("lib/games/word_search/word_search_models.dart", "05-19", "10, 12"),
    ("lib/games/word_search/word_search_service.dart", "05-19", "10, 12"),
    ("lib/games/rock_paper_scissors/screens/rps_intro_screen.dart", "05-19", "9"),
    ("lib/games/game_hub_screen.dart", "05-19, 05-20", "9, 16, 20, 21"),
    ("lib/screens/chats_screen.dart", "05-19, 05-20", "17, 24, 25"),
    ("lib/screens/main_navigation.dart", "05-20", "22"),
    ("lib/screens/home_screen.dart", "05-20", "23"),
    ("lib/services/matching_service.dart", "05-20", "22 (upsert fix)"),
    ("Supabase DB (migrations)", "05-18, 05-19", "3, 5, 12"),
]
for fname, session, fixes in file_rows:
    row = table2.add_row().cells
    row[0].text = fname
    row[1].text = session
    row[2].text = fixes

doc.add_paragraph()

# ── Screenshot Guide ───────────────────────────────────────────────────────────
add_heading(doc, "7. Screenshot Guide", level=1)
doc.add_paragraph(
    "Below is a consolidated list of recommended screenshots to include in this report. "
    "Each entry describes what to show and which code location generates that UI."
)

screenshots = [
    (
        "1. Emoji Charades — waitSolve phase",
        "lib/games/emoji_charades/emoji_charades_game_screen.dart",
        "_buildWaitSolve() or _goTo(_Phase.waitSolve) — show the 'Waiting for partner to solve' screen",
    ),
    (
        "2. Chat — expired game invitation card",
        "lib/screens/chat_conversation_screen.dart",
        "_buildGameChallengeBubble() — the dark-grey expired card with 'Game session expired' text",
    ),
    (
        "3. RPS — both players submitted, reveal screen",
        "lib/games/rock_paper_scissors/screens/rps_reveal_screen.dart",
        "Show both moves visible after the RLS fix",
    ),
    (
        "4. Chat — RPS game card showing 'Game completed ✓'",
        "lib/screens/chat_conversation_screen.dart",
        "_buildGameChallengeBubble() after onChatUnlocked refresh",
    ),
    (
        "5. Login screen (no black screen after Google Sign-In back)",
        "lib/screens/login_screen.dart",
        "Navigate back from Google Sign-In — Impeller disabled ensures no black screen",
    ),
    (
        "6. RPS intro/rules screen",
        "lib/games/rock_paper_scissors/screens/rps_intro_screen.dart",
        "First screen shown before RPS gameplay",
    ),
    (
        "7. Word Search welcome screen",
        "lib/games/word_search/word_search_supabase_wrapper.dart",
        "_buildWelcome() — skipWelcome now false",
    ),
    (
        "8. Word Search result screen (win/loss/draw)",
        "lib/games/word_search/word_search_game_screen.dart",
        "_buildDone() — shows win/lose/draw label based on isSkipped flags",
    ),
    (
        "9. Word Search wait screen ('Partner's puzzle is ready!')",
        "lib/games/word_search/word_search_game_screen.dart",
        "_buildWait() — shows 'Solve [Partner]'s Puzzle' button when partner.submitted = true",
    ),
    (
        "10. Emoji Charades waiting screen ('Solve Partner's Puzzle' button)",
        "lib/games/emoji_charades/emoji_charades_game_screen.dart",
        "_buildWaiting() — button shown when _partnerSubmitted = true",
    ),
    (
        "11. Chat '+' picker (no controller tile)",
        "lib/screens/chat_conversation_screen.dart",
        "_showMediaPicker() — game controller tile removed",
    ),
    (
        "12. GameHub — challenge card in chat",
        "lib/screens/chat_conversation_screen.dart",
        "_buildGameChallengeBubble() — card for both sender and receiver",
    ),
    (
        "13. GameHub — title/subtitle gradient",
        "lib/games/game_hub_screen.dart",
        "ShaderMask on 'Pick a game to play' and subtitle text",
    ),
    (
        "14. MainNavigation — 'It's a Match!' banner (User A)",
        "lib/screens/main_navigation.dart",
        "Notification banner shown via _pollForNewMatches() within 15 seconds of match",
    ),
    (
        "15. HomeScreen — 'It's a Match!' popup with updated button",
        "lib/screens/home_screen.dart",
        "Button text 'PLAY GAME TO UNLOCK CHAT'",
    ),
    (
        "16. Chats screen — Pending Games circles + Messages section",
        "lib/screens/chats_screen.dart",
        "Color-coded borders (grey/green/red/purple) + split layout",
    ),
    (
        "17. Chats screen header (no controller button)",
        "lib/screens/chats_screen.dart",
        "Header — sports_esports icon removed",
    ),
    (
        "18. Chat — 'Rejoin Game' button on active card",
        "lib/screens/chat_conversation_screen.dart",
        "_buildGameChallengeBubble() — purple 'Rejoin Game' button when canRejoin = true",
    ),
]

for title_text, file_ref, detail in screenshots:
    p = doc.add_paragraph()
    run_title = p.add_run(title_text + "\n")
    run_title.bold = True
    run_file = p.add_run("  Code: " + file_ref + "\n")
    run_file.italic = True
    run_file.font.size = Pt(9)
    p.add_run("  " + detail)
    p.paragraph_format.space_after = Pt(6)

# Save
out_path = r"c:\group prj\matchxp\MatchXP_Session_Report_2026-05-18_to_20.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
